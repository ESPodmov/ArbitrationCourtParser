import docx
from utils import add_hyperlink


def create_or_edit_docx_from_list(path, my_list: list, format_line, link_dict: dict = None):
    document = docx.Document()

    for dictionary in my_list:
        dict_to_paragraphs_docx(document, dictionary, format_line, link_dict)

    document.save(path)


def dict_to_paragraphs_docx(document, dictionary: dict, format_line, link_dict: dict):
    paragraphs_list = format_line.split("\n")
    link_dict_keys = list(link_dict.keys())

    for paragraph in paragraphs_list:
        doc_paragraph = document.add_paragraph()
        flag, current_list = paragraph_to_list_with_links(paragraph, link_dict)
        if flag:
            for elem in current_list:
                if elem in link_dict_keys:
                    link = link_dict[elem]
                    add_hyperlink(doc_paragraph, link, link)
                else:
                    doc_paragraph.add_run(elem.format_map(dictionary))
        else:
            doc_paragraph.add_run(current_list[0].format_map(dictionary))


def paragraph_to_list_with_links(paragraph: str, link_dict: dict):
    split_line = [paragraph]
    if link_dict is not None:
        for key in list(link_dict.keys()):
            main_split = split_line.copy()
            for i in range(len(split_line)):
                if "{" + key + "}" in split_line[i]:
                    current_split = split_line[i].split("{" + key + "}")
                    for j in range(len(current_split) - 1):
                        current_split.insert(j * 2 + 1, key)
                    current_split = list(filter(lambda x: x != "", current_split))
                    index = -len(main_split) + i
                    main_split[index:index] = current_split
                    main_split.pop(index)
            split_line = main_split.copy()
            del main_split
        return not len(split_line) == 1, split_line
    else:
        return False, split_line
